# mlapp/ml/docx_processor.py
import os
from docx import Document
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from django.conf import settings
import uuid

def extract_text_from_docx(docx_file, output_file=None):
    """
    Extract text from a DOCX file and save it to a text file.
    
    Args:
        docx_file: Can be either a file path string or a Django UploadedFile object
        output_file: Path where the extracted text will be saved. If None, a default path will be generated.
    
    Returns:
        str: Path to the saved text file
    """
    # If output_file is not provided, generate a default filename
    if output_file is None:
        # Create a directory for extracted texts if it doesn't exist
        extract_dir = os.path.join(settings.MEDIA_ROOT, 'extracted_texts')
        os.makedirs(extract_dir, exist_ok=True)
        
        # Generate a unique filename
        filename = f"docx_extract_{uuid.uuid4().hex}.txt"
        output_file = os.path.join(extract_dir, filename)
    
    # If input is a Django UploadedFile or InMemoryUploadedFile
    if hasattr(docx_file, 'read'):
        # Save the uploaded file temporarily
        temp_path = default_storage.save('mlapp/temp_docs/temp.docx', ContentFile(docx_file.read()))
        try:
            # Reset file pointer for future use
            docx_file.seek(0)
            # Open the document from the temp file
            doc = Document(default_storage.path(temp_path))
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Write extracted text to file
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(text)
            
            return output_file
        finally:
            # Clean up the temp file
            if default_storage.exists(temp_path):
                default_storage.delete(temp_path)
    
    # If input is a file path
    else:
        # Assume docx_file is a path string
        if not os.path.exists(docx_file):
            raise FileNotFoundError(f"File not found: {docx_file}")
        
        doc = Document(docx_file)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # Write extracted text to file
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(text)
        
        return output_file


# -------------------------------------------------------------------------
# USAGE EXAMPLES
# -------------------------------------------------------------------------

# Example 1: Using in a Django view with an uploaded file
"""
# In any Django view:
def process_docx_view(request):
    from django.http import JsonResponse
    import os
    
    if request.method == 'POST' and request.FILES.get('document'):
        docx_file = request.FILES['document']
        
        # Check if the file is a docx
        if not docx_file.name.endswith('.docx'):
            return JsonResponse({'error': 'Please upload a DOCX file'}, status=400)
        
        try:
            # Extract and save to a text file
            output_path = extract_text_from_docx(docx_file)
            
            # Get the relative path for client-side use
            relative_path = os.path.relpath(output_path, settings.MEDIA_ROOT)
            url_path = os.path.join(settings.MEDIA_URL, relative_path)
            
            return JsonResponse({
                'success': True, 
                'file_path': output_path,
                'download_url': url_path
            })
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    
    return JsonResponse({'error': 'Invalid request'}, status=400)
"""

# Example 2: Using with a file path and custom output
"""
# Anywhere in your Django project:
def process_existing_docx(file_path, output_path=None):
    try:
        result_path = extract_text_from_docx(file_path, output_path)
        print(f"Text extracted and saved to: {result_path}")
        return result_path
    except FileNotFoundError:
        print(f"The file {file_path} does not exist")
        return None
    except Exception as e:
        print(f"Error processing file: {e}")
        return None

# Example call
# output = process_existing_docx('/path/to/your/document.docx', '/path/to/output/text.txt')
"""

# Example 3: Using in a Django model method
"""
# In a Django model:
class Document(models.Model):
    title = models.CharField(max_length=255)
    file = models.FileField(upload_to='documents/')
    extracted_text_file = models.FileField(upload_to='extracted_texts/', blank=True)
    
    def extract_content(self):
        if self.file and self.file.name.endswith('.docx'):
            # Generate output path
            filename = f"extract_{os.path.basename(self.file.name)}.txt"
            output_path = os.path.join(settings.MEDIA_ROOT, 'extracted_texts', filename)
            
            # Extract text to file
            result_path = extract_text_from_docx(self.file.path, output_path)
            
            # Update model with the path
            relative_path = os.path.relpath(result_path, settings.MEDIA_ROOT)
            self.extracted_text_file = relative_path
            self.save()
            return True
        return False
"""

# Example 4: Using from another app
"""
# In another_app/views.py:
from mlapp.ml.docx_processor import extract_text_from_docx

def external_app_process(request):
    # Similar to Example 1
    pass
"""